"""
Template mapper for Omar Zúñiga's dictamen template.
Maps analysis data to the specific fields in the Word template.
"""

from docx import Document
from typing import Dict, Any
from datetime import datetime


class TemplateMapper:
    """Maps analysis data to Word template fields."""
    
    def __init__(self, doc: Document):
        """Initialize with a Document object."""
        self.doc = doc
    
    def map_data(self, dictamen_data: Dict[str, Any]):
        """
        Map dictamen data to template fields.
        
        Args:
            dictamen_data: Dictionary with analysis data
        """
        # Map data to first table (DATOS DE LA INSTALACIÓN)
        if len(self.doc.tables) >= 1:
            self._map_installation_data(self.doc.tables[0], dictamen_data)
        
        # Map data to second table (Ubicación)
        if len(self.doc.tables) >= 2:
            self._map_location_data(self.doc.tables[1], dictamen_data)
        
        # Replace analysis section
        self._replace_analysis_section(dictamen_data)
        
        # Replace conclusion
        self._replace_conclusion(dictamen_data)
    
    def _map_installation_data(self, table, data: Dict[str, Any]):
        """Map data to installation data table."""
        # Table structure:
        # Row 0: Headers (DATO | INFORMACIÓN)
        # Row 1: Folio No.
        # Row 2: Fecha del Dictamen
        # Row 3: Nombre de la Instalación
        # Row 4: RFC
        # Row 5-8: Other data
        
        try:
            # Folio (Row 1, Col 1)
            if len(table.rows) > 1:
                table.rows[1].cells[1].text = data.get('folio', 'DICT-AUTO-001')
            
            # Fecha (Row 2, Col 1)
            if len(table.rows) > 2:
                fecha = data.get('fecha', datetime.now().strftime('%d de %B de %Y'))
                table.rows[2].cells[1].text = fecha
            
            # Nombre de la Instalación (Row 3, Col 1)
            if len(table.rows) > 3:
                tipo = data.get('tipo_instalacion', 'Instalación Eléctrica')
                table.rows[3].cells[1].text = f"{tipo}"
            
            # RFC (Row 4, Col 1) - optional
            if len(table.rows) > 4:
                rfc = data.get('rfc', data.get('solicitante', 'N/A'))
                table.rows[4].cells[1].text = rfc
                
        except Exception as e:
            print(f"Warning: Error mapping installation data: {e}")
    
    def _map_location_data(self, table, data: Dict[str, Any]):
        """Map data to location table."""
        # Table structure:
        # Row 1: Ubicación
        # Row 2: Municipio
        # Row 3: Estado
        # Row 4: Código Postal
        
        try:
            ubicacion_data = data.get('ubicacion', '')
            
            # Parse ubicacion if it's a string
            if isinstance(ubicacion_data, str):
                # Ubicación (Row 1, Col 1)
                if len(table.rows) > 1:
                    table.rows[1].cells[1].text = ubicacion_data
            elif isinstance(ubicacion_data, dict):
                # Ubicación detallada
                if len(table.rows) > 1:
                    table.rows[1].cells[1].text = ubicacion_data.get('direccion', 'N/A')
                if len(table.rows) > 2:
                    table.rows[2].cells[1].text = ubicacion_data.get('municipio', 'N/A')
                if len(table.rows) > 3:
                    table.rows[3].cells[1].text = ubicacion_data.get('estado', 'N/A')
                if len(table.rows) > 4:
                    table.rows[4].cells[1].text = ubicacion_data.get('cp', 'N/A')
                    
        except Exception as e:
            print(f"Warning: Error mapping location data: {e}")
    
    def _replace_analysis_section(self, data: Dict[str, Any]):
        """Replace the analysis section with actual findings."""
        # Find the "ANÁLISIS VISUAL vs. NOM-001-SEDE-2012" section
        analysis_start_idx = None
        conclusion_idx = None
        
        for i, para in enumerate(self.doc.paragraphs):
            if 'ANÁLISIS VISUAL' in para.text.upper():
                analysis_start_idx = i
            if 'CONCLUSIÓN' in para.text.upper():
                conclusion_idx = i
                break
        
        if analysis_start_idx is not None and conclusion_idx is not None:
            # Remove old analysis paragraphs
            for i in range(conclusion_idx - 1, analysis_start_idx, -1):
                self._delete_paragraph(self.doc.paragraphs[i])
            
            # Insert new analysis
            insert_point = analysis_start_idx + 1
            
            # Add summary paragraph
            summary = data.get('observaciones', 'Análisis en proceso...')
            self._insert_paragraph_at(insert_point, summary)
            insert_point += 1
            
            # Add detailed findings
            non_conformities = data.get('no_conformidades_detalladas', [])
            if non_conformities:
                for nc in non_conformities:
                    # Add section title
                    self._insert_paragraph_at(insert_point, f"\n{nc.get('titulo', 'Observación')}", bold=True)
                    insert_point += 1
                    
                    # Add description
                    self._insert_paragraph_at(insert_point, nc.get('descripcion', ''))
                    insert_point += 1
                    
                    # Add norm reference
                    if nc.get('norma'):
                        self._insert_paragraph_at(insert_point, f"Lo que dice la NOM:\n{nc.get('norma')}")
                        insert_point += 1
                    
                    # Add verdict
                    if nc.get('veredicto'):
                        self._insert_paragraph_at(insert_point, f"Veredicto: {nc.get('veredicto')}")
                        insert_point += 1
    
    def _replace_conclusion(self, data: Dict[str, Any]):
        """Replace conclusion section."""
        for para in self.doc.paragraphs:
            if 'CONCLUSIÓN' in para.text.upper():
                # Find next paragraph and replace
                idx = self.doc.paragraphs.index(para)
                if idx + 1 < len(self.doc.paragraphs):
                    conclusion_text = f"{data.get('dictamen_final', 'PENDIENTE')}\n\n{data.get('justificacion', '')}"
                    self.doc.paragraphs[idx + 1].text = conclusion_text
                break
    
    def _delete_paragraph(self, paragraph):
        """Delete a paragraph from document."""
        p = paragraph._element
        p.getparent().remove(p)
    
    def _insert_paragraph_at(self, index: int, text: str, bold: bool = False):
        """Insert a paragraph at specific index."""
        if index < len(self.doc.paragraphs):
            p = self.doc.paragraphs[index]._element
            new_p = self.doc.add_paragraph(text)._element
            p.addprevious(new_p)
            
            if bold and len(self.doc.paragraphs) > index:
                for run in self.doc.paragraphs[index].runs:
                    run.bold = True
        else:
            p = self.doc.add_paragraph(text)
            if bold:
                for run in p.runs:
                    run.bold = True


def map_data_to_template(doc: Document, dictamen_data: Dict[str, Any]):
    """
    Convenience function to map data to template.
    
    Args:
        doc: Document object
        dictamen_data: Data to map
    """
    mapper = TemplateMapper(doc)
    mapper.map_data(dictamen_data)

"""
Word Document Generator for Electrical Inspection Dictamen
Generates .docx documents matching the PDF format for easy editing and copy/paste.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Set


class WordGenerator:
    """Generate Word dictamen documents."""
    
    def __init__(self, output_dir: str = "data/generated"):
        """Initialize Word generator."""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def generate_dictamen(self, data: Dict[str, Any], image_path: str = None) -> str:
        """
        Generate Word dictamen from analysis data.
        
        Args:
            data: Dictionary containing analysis results
            image_path: Optional path to the analyzed image
            
        Returns:
            Path to generated Word file
        """
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Dictamen_AUTO-{int(datetime.now().timestamp() * 1000)}_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Title
        title = doc.add_paragraph()
        title_run = title.add_run("Dictamen Técnico de Instalación Eléctrica (Basado en NOM-001-SEDE-2012)")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 82, 130)  # #2C5282
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()
        
        # Metadata
        now = datetime.now()
        meses_es = {
            1: 'enero', 2: 'febrero', 3: 'marzo', 4: 'abril',
            5: 'mayo', 6: 'junio', 7: 'julio', 8: 'agosto',
            9: 'septiembre', 10: 'octubre', 11: 'noviembre', 12: 'diciembre'
        }
        fecha = f"{now.day} de {meses_es[now.month]} de {now.year}"
        
        p = doc.add_paragraph()
        p.add_run("Fecha del Dictamen: ").bold = True
        p.add_run(fecha)
        
        p = doc.add_paragraph()
        p.add_run("Referencia: ").bold = True
        p.add_run("Análisis de Imagen(es) de Instalación Eléctrica")
        
        p = doc.add_paragraph()
        p.add_run("Normativa Aplicable: ").bold = True
        p.add_run("Norma Oficial Mexicana NOM-001-SEDE-2012, Instalaciones Eléctricas (Utilización). (Se reconoce que la NOM-001-SEDE-2012 se basa en el National Electrical Code, NFPA 70, y las referencias numéricas proporcionadas en las imágenes corresponden a artículos de dicho código, los cuales están integrados en la estructura de la NOM).")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        
        # 1. Introducción
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("1. Introducción")
        h1_run.font.size = Pt(13)
        h1_run.font.bold = True
        h1_run.font.color.rgb = RGBColor(44, 82, 130)
        
        intro = doc.add_paragraph()
        intro.add_run("El presente dictamen técnico tiene como objetivo analizar la(s) imagen(es) proporcionada(s) de una instalación eléctrica, con especial atención a la distribución de conductores dentro de un tablero de distribución o centro de carga. Se evaluará el cumplimiento de los principios fundamentales de seguridad, diseño, selección y construcción establecidos en la NOM-001-SEDE-2012, identificando aspectos conformes, no conformes y aquellos que requieren verificación adicional, así como proporcionando recomendaciones para subsanar deficiencias.")
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        
        # 2. Análisis Detallado
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("2. Análisis Detallado de la Instalación")
        h1_run.font.size = Pt(13)
        h1_run.font.bold = True
        h1_run.font.color.rgb = RGBColor(44, 82, 130)
        
        p = doc.add_paragraph()
        p.add_run("A continuación, se presenta un análisis de los elementos visibles en la imagen, en relación con las referencias normativas señaladas y la NOM-001-SEDE-2012:")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Insert image here (Section 2, before 2.1)
        if image_path and Path(image_path).exists():
            try:
                doc.add_paragraph()
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(image_path, width=Inches(5.5))
                doc.add_paragraph()
            except Exception as e:
                print(f"Error inserting image into Word: {e}")
        
        doc.add_paragraph()
        
        # Get data
        non_conformities = data.get('non_conformities', [])
        conformities = data.get('conformities', [])
        
        # 2.1 Aspectos que cumplen
        h2 = doc.add_paragraph()
        h2_run = h2.add_run("2.1. Aspectos que cumplen con la normativa (✓)")
        h2_run.font.size = Pt(11)
        h2_run.font.bold = True
        
        if conformities:
            for conf in conformities[:5]:
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"✓ {conf}")
                run.bold = True
        else:
            doc.add_paragraph("• No se identificaron aspectos conformes específicos en el análisis visual.")
        
        doc.add_paragraph()
        
        # 2.2 Aspectos que NO cumplen
        h2 = doc.add_paragraph()
        h2_run = h2.add_run("2.2. Aspectos que NO cumplen o presentan riesgos (✗)")
        h2_run.font.size = Pt(11)
        h2_run.font.bold = True
        
        if non_conformities:
            for nc in non_conformities:
                desc = nc.get('description', 'Sin descripción')
                article = nc.get('article', 'Sin referencia')
                severity = nc.get('severity', 'medium')
                
                # Title
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"✗ {desc}")
                run.bold = True
                
                # Observación
                p = doc.add_paragraph()
                p.add_run("Observación: ").bold = True
                p.add_run(desc)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Riesgo
                risk_map = {
                    'high': 'Severo. Riesgo inminente de daño al aislamiento de los conductores por abrasión o corte, lo que podría provocar cortocircuitos, fallas a tierra, arcos eléctricos e incluso incendios.',
                    'medium': 'Alto. Puede generar sobrecalentamiento, fallas en la protección y riesgo de incendio.',
                    'low': 'Moderado. Puede afectar la seguridad y eficiencia de la instalación a largo plazo.'
                }
                p = doc.add_paragraph()
                p.add_run("Riesgo: ").bold = True
                p.add_run(risk_map.get(severity, risk_map['medium']))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Normativa Aplicable
                if article and article != 'Sin referencia':
                    p = doc.add_paragraph()
                    p.add_run("Normativa Aplicable: ").bold = True
                    run = p.add_run(f"NOM-001-SEDE-2012, Artículo {article}")
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                
                doc.add_paragraph()
        else:
            doc.add_paragraph("• No se identificaron no conformidades en el análisis visual.")
        
        doc.add_paragraph()
        
        # 3. Recomendaciones
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("3. Recomendaciones Específicas de Corrección")
        h1_run.font.size = Pt(13)
        h1_run.font.bold = True
        h1_run.font.color.rgb = RGBColor(44, 82, 130)
        
        recommendations = [
            ("Protección en Aberturas Metálicas", "Instalar de inmediato pasacables, bujes o anillos aprobados en todas las aberturas metálicas por donde ingresan los conductores al tablero, asegurando que cubran completamente los bordes metálicos."),
            ("Manejo de Conductores y Disipación de Calor", "Deshacer el agrupamiento excesivo de conductores o, en su defecto, aplicar los factores de ajuste de ampacidad correspondientes según la Tabla 310.15(B)(3)(a) de la NOM-001-SEDE-2012 para asegurar que los conductores no se sobrecalienten."),
            ("Organización del Cableado", "Reorganizar el cableado dentro del tablero para un tendido más limpio y ordenado, utilizando cinchos o sujetacables de forma no restrictiva para mantener los conductores en su lugar sin apretarlos excesivamente.")
        ]
        
        for i, (title, desc) in enumerate(recommendations, 1):
            p = doc.add_paragraph()
            p.add_run(f"{i}. {title}: ").bold = True
            p.add_run(desc)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        
        # 4. Conclusión
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("4. Conclusión")
        h1_run.font.size = Pt(13)
        h1_run.font.bold = True
        h1_run.font.color.rgb = RGBColor(44, 82, 130)
        
        classification = data.get('classification', {})
        justification = classification.get('justification', '')
        
        if justification:
            p = doc.add_paragraph()
            p.add_run(justification)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            conclusion_text = "La instalación eléctrica analizada presenta deficiencias significativas en cuanto a la protección mecánica de los conductores y organización del cableado. Estas no conformidades representan riesgos serios para la seguridad de las personas y la propiedad, y deben ser corregidas de manera prioritaria para asegurar el cumplimiento con la NOM-001-SEDE-2012. Se recomienda encarecidamente la intervención de personal calificado para realizar las modificaciones necesarias y garantizar la seguridad y fiabilidad de la instalación."
            p = doc.add_paragraph()
            p.add_run(conclusion_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Elaborado por
        inspector_name = data.get('inspector_name', '[ Tu Nombre ]')
        p = doc.add_paragraph()
        p.add_run("Elaborado por: ").bold = True
        p.add_run(inspector_name)
        
        doc.add_paragraph()
        
        # Referencias de NOMs
        p = doc.add_paragraph()
        p.add_run("Referencias de NOMs:")
        
        # Extract unique articles
        articles_set: Set[str] = set()
        for nc in non_conformities:
            article = nc.get('article')
            if article and article != 'Sin referencia':
                articles_set.add(article)
        
        # Generate references
        if articles_set:
            for article in sorted(articles_set):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"NOM-001-SEDE-2012.pdf (Artículo {article})")
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        else:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run("NOM-001-SEDE-2012.pdf (Referencia general)")
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        
        # Save document
        doc.save(str(filepath))
        
        print(f"✓ Word generated: {filepath}")
        return str(filepath)

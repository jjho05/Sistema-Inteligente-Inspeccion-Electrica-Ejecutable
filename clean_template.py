"""
Script to clean Omar Zúñiga's contact information from Word template.
"""

from docx import Document
from pathlib import Path

def clean_template():
    """Remove Omar Zúñiga's contact info from template."""
    
    template_path = Path("data/templates/DictamenElectrico_OmarZuniga.docx")
    
    if not template_path.exists():
        print(f"❌ Template not found: {template_path}")
        return
    
    print(f"📄 Loading template: {template_path}")
    doc = Document(str(template_path))
    
    # Text to remove
    remove_texts = [
        "M.E. Omar Zúñiga Núñez",
        "UVSEIE 717",
        "Domicilio: Av. Ayuntamiento No.2400, Int. 1er. Piso, Col. Lauro Aguirre C.P. 89140 Tampico, Tamaulipas",
        "Teléfono: 01-833-213-2664",
        "Celular: 833 1 85 36 20",
        "Correo electrónico: uv_ingoz@hotmail.com",
        "Omar Zúñiga",
        "uv_ingoz@hotmail.com",
        "833 1 85 36 20",
        "01-833-213-2664"
    ]
    
    removed_count = 0
    
    # Clean paragraphs
    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        for remove_text in remove_texts:
            if remove_text.lower() in original_text.lower():
                # Clear the paragraph
                for run in paragraph.runs:
                    run.text = ""
                removed_count += 1
                print(f"  ✓ Removed: {original_text[:50]}...")
                break
    
    # Clean tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    original_text = paragraph.text
                    for remove_text in remove_texts:
                        if remove_text.lower() in original_text.lower():
                            for run in paragraph.runs:
                                run.text = ""
                            removed_count += 1
                            print(f"  ✓ Removed from table: {original_text[:50]}...")
                            break
    
    # Save cleaned template
    output_path = Path("data/templates/DictamenElectrico_Cleaned.docx")
    doc.save(str(output_path))
    
    print(f"\n✅ Template cleaned!")
    print(f"   Removed {removed_count} instances")
    print(f"   Saved to: {output_path}")
    print(f"\n💡 Update template_mapper.py to use: DictamenElectrico_Cleaned.docx")

if __name__ == "__main__":
    clean_template()

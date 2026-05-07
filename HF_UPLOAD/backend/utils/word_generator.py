"""
Word Document Generator for Electrical Inspection Dictamen
Generates .docx documents matching the PDF format for easy editing and copy/paste.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Set, List


class WordGenerator:
    """Generate Word dictamen documents."""
    
    def __init__(self, output_dir: str = "data/generated"):
        """Initialize Word generator."""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def generate_dictamen(self, data: Dict[str, Any], image_paths: List[str] = None, image_path: str = None, language: str = 'es') -> str:
        """Generate Word dictamen. language='en' produces an English report."""
        en = (language == 'en')
        if image_path and not image_paths:
            image_paths = [image_path]
        elif not image_paths:
            image_paths = []
            
        # Generate filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"Dictamen_AUTO-{int(datetime.now().timestamp() * 1000)}_{timestamp}.docx"
        filepath = self.output_dir / filename
        
        # Create document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Title
        doc_title = "Technical Report – Electrical Installation (NOM-001-SEDE-2012 / NEC)" if en else "Dictamen Técnico de Instalación Eléctrica (Basado en NOM-001-SEDE-2012)"
        title = doc.add_paragraph()
        title_run = title.add_run(doc_title)
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(44, 82, 130)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph()

        # Metadata
        now = datetime.now()
        if en:
            fecha = now.strftime('%B %d, %Y')
            p = doc.add_paragraph(); p.add_run("Report Date: ").bold = True; p.add_run(fecha)
            p = doc.add_paragraph(); p.add_run("Reference: ").bold = True; p.add_run("AI Analysis of Electrical Installation Image(s)")
            p = doc.add_paragraph(); p.add_run("Applicable Standard: ").bold = True
            p.add_run("NOM-001-SEDE-2012 / National Electrical Code (NFPA 70). Numerical article references correspond to NEC articles integrated into the NOM structure.")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            meses_es = {1:'enero',2:'febrero',3:'marzo',4:'abril',5:'mayo',6:'junio',7:'julio',8:'agosto',9:'septiembre',10:'octubre',11:'noviembre',12:'diciembre'}
            fecha = f"{now.day} de {meses_es[now.month]} de {now.year}"
            p = doc.add_paragraph(); p.add_run("Fecha del Dictamen: ").bold = True; p.add_run(fecha)
            p = doc.add_paragraph(); p.add_run("Referencia: ").bold = True; p.add_run("Análisis de Imagen(es) de Instalación Eléctrica")
            p = doc.add_paragraph(); p.add_run("Normativa Aplicable: ").bold = True
            p.add_run("Norma Oficial Mexicana NOM-001-SEDE-2012, Instalaciones Eléctricas (Utilización). (Se reconoce que la NOM-001-SEDE-2012 se basa en el National Electrical Code, NFPA 70).")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

        # 1. Introduction
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("1. Introduction" if en else "1. Introducción")
        h1_run.font.size = Pt(13); h1_run.font.bold = True; h1_run.font.color.rgb = RGBColor(44, 82, 130)
        intro = doc.add_paragraph()
        intro.add_run("This technical report analyzes the provided electrical installation image(s), focusing on conductor distribution and safety compliance with NOM-001-SEDE-2012 / NEC standards."
                      if en else
                      "El presente dictamen técnico analiza la(s) imagen(es) proporcionada(s) de una instalación eléctrica, evaluando el cumplimiento con la NOM-001-SEDE-2012.")
        intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

        # 2. Detailed Analysis
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("2. Detailed Analysis" if en else "2. Análisis Detallado de la Instalación")
        h1_run.font.size = Pt(13); h1_run.font.bold = True; h1_run.font.color.rgb = RGBColor(44, 82, 130)
        p = doc.add_paragraph()
        p.add_run("The following elements were identified in the image(s):" if en else "A continuación, se presenta un análisis de los elementos visibles en la(s) imagen(es):")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Insert images (Sequential)
        if image_paths:
            doc.add_paragraph()
            for i, img_path in enumerate(image_paths):
                if Path(img_path).exists():
                    try:
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = img_para.add_run()
                        run.add_picture(img_path, width=Inches(5.5))
                        
                        # Add caption
                        caption = doc.add_paragraph(f"Figura {i+1}: Vista de la instalación analizada")
                        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        caption.style = 'Caption' if 'Caption' in doc.styles else 'Normal'
                        
                        doc.add_paragraph()
                    except Exception as e:
                        print(f"Error inserting image into Word: {e}")
        
        # Get data
        non_conformities = data.get('non_conformities', [])
        conformities = data.get('conformities', [])

        # 2.1 Conforming aspects
        h2 = doc.add_paragraph()
        h2_run = h2.add_run("2.1. Conforming Aspects (✓)" if en else "2.1. Aspectos que cumplen con la normativa (✓)")
        h2_run.font.size = Pt(11); h2_run.font.bold = True
        if conformities:
            for conf in conformities[:5]:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(f"✓ {conf}").bold = True
        else:
            doc.add_paragraph("• No specific conforming aspects were identified." if en else "• No se identificaron aspectos conformes específicos en el análisis visual.")
        doc.add_paragraph()

        # 2.2 Non-conforming aspects
        h2 = doc.add_paragraph()
        h2_run = h2.add_run("2.2. Non-Conforming Aspects / Risks (✗)" if en else "2.2. Aspectos que NO cumplen o presentan riesgos (✗)")
        h2_run.font.size = Pt(11); h2_run.font.bold = True
        if non_conformities:
            risk_map_en = {'high': 'Severe. Imminent risk of conductor insulation damage.', 'medium': 'High. May cause overheating and fire risk.', 'low': 'Moderate. May affect safety over time.'}
            risk_map_es = {'high': 'Severo. Riesgo inminente de daño al aislamiento.', 'medium': 'Alto. Puede generar sobrecalentamiento y riesgo de incendio.', 'low': 'Moderado. Puede afectar la seguridad a largo plazo.'}
            risk_map = risk_map_en if en else risk_map_es
            for nc in non_conformities:
                desc = nc.get('description', 'No description' if en else 'Sin descripción')
                article = nc.get('article', 'No reference' if en else 'Sin referencia')
                severity = nc.get('severity', 'medium')
                p = doc.add_paragraph(style='List Bullet'); p.add_run(f"✗ {desc}").bold = True
                obs_lbl = "Observation: " if en else "Observación: "
                p = doc.add_paragraph(); p.add_run(obs_lbl).bold = True; p.add_run(desc); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                risk_lbl = "Risk: " if en else "Riesgo: "
                p = doc.add_paragraph(); p.add_run(risk_lbl).bold = True; p.add_run(risk_map.get(severity, risk_map['medium'])); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if article and article not in ('Sin referencia', 'No reference'):
                    norm_lbl = "Applicable Standard: " if en else "Normativa Aplicable: "
                    norm_ref = f"NOM-001-SEDE-2012 / NEC, Article {article}" if en else f"NOM-001-SEDE-2012, Artículo {article}"
                    p = doc.add_paragraph(); p.add_run(norm_lbl).bold = True
                    run = p.add_run(norm_ref); run.font.color.rgb = RGBColor(255, 0, 0); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                doc.add_paragraph()
        else:
            doc.add_paragraph("• No non-conformities were detected." if en else "• No se identificaron no conformidades en el análisis visual.")
        doc.add_paragraph()

        # 3. Recommendations
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("3. Specific Correction Recommendations" if en else "3. Recomendaciones Específicas de Corrección")
        h1_run.font.size = Pt(13); h1_run.font.bold = True; h1_run.font.color.rgb = RGBColor(44, 82, 130)
        if en:
            recommendations = [
                ("Conductor Protection at Metal Openings", "Immediately install approved cable grommets or bushings at all metal openings."),
                ("Conductor Bundling & Heat Dissipation", "Reduce bundling or apply ampacity adjustment factors per Table 310.15(B)(3)(a)."),
                ("Cable Organization", "Reorganize wiring for a cleaner layout using non-restrictive cable ties."),
            ]
        else:
            recommendations = [
                ("Protección en Aberturas Metálicas", "Instalar de inmediato pasacables, bujes o anillos aprobados en todas las aberturas metálicas."),
                ("Manejo de Conductores y Disipación de Calor", "Deshacer el agrupamiento excesivo o aplicar factores de ajuste de ampacidad según Tabla 310.15(B)(3)(a)."),
                ("Organización del Cableado", "Reorganizar el cableado dentro del tablero usando cinchos o sujetacables de forma no restrictiva."),
            ]
        for i, (title, desc) in enumerate(recommendations, 1):
            p = doc.add_paragraph(); p.add_run(f"{i}. {title}: ").bold = True; p.add_run(desc); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

        # 4. Conclusion
        h1 = doc.add_paragraph()
        h1_run = h1.add_run("4. Conclusion" if en else "4. Conclusión")
        h1_run.font.size = Pt(13); h1_run.font.bold = True; h1_run.font.color.rgb = RGBColor(44, 82, 130)
        
        classification = data.get('classification', {})
        justification = classification.get('justification', '')
        
        if justification:
            p = doc.add_paragraph()
            p.add_run(justification)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            conclusion_text = "La instalación eléctrica analizada presenta deficiencias significativas en cuanto a la protección mecánica de los conductores y organización del cableado. Estas no conformidades representan riesgos serios para la seguridad de las personas y la propiedad, y deben ser corregidas de manera prioritaria para asegurar el cumplimiento con la NOM-001-SEDE-2012. Se recomienda encarecidamente la intervención de personal calificado para realizar las modificaciones necesarias y garantizar la seguridad y fiabilidad de la instalación."
            p = doc.add_paragraph()
            p.add_run(conclusion_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        inspector_name = data.get('inspector_name', '[ Inspector ]')
        p = doc.add_paragraph()
        p.add_run("Prepared by: " if en else "Elaborado por: ").bold = True
        p.add_run(inspector_name)
        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("Standards References:" if en else "Referencias de NOMs:")
        articles_set: Set[str] = set()
        for nc in non_conformities:
            article = nc.get('article')
            if article and article not in ('Sin referencia', 'No reference'):
                articles_set.add(article)
        if articles_set:
            for article in sorted(articles_set):
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(f"NOM-001-SEDE-2012 / NEC (Article {article})" if en else f"NOM-001-SEDE-2012.pdf (Artículo {article})")
                run.font.name = 'Courier New'; run.font.size = Pt(9)
        else:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run("NOM-001-SEDE-2012 / NEC (General Reference)" if en else "NOM-001-SEDE-2012.pdf (Referencia general)")
            run.font.name = 'Courier New'; run.font.size = Pt(9)
        
        # Save document
        doc.save(str(filepath))
        
        print(f"✓ Word generated: {filepath}")
        return str(filepath)
